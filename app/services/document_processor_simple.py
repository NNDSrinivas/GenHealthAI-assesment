import os
import re
import logging
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    
    def __init__(self, tesseract_path: Optional[str] = None, poppler_path: Optional[str] = None):
        self.tesseract_path = tesseract_path
        self.poppler_path = poppler_path
        # Only support text-based formats for now
        self.supported_formats = {'.docx', '.txt'}
        
        # Patient data extraction patterns
        self.name_patterns = [
            r'patient\s*name[:\s]+([a-zA-Z\s,]+)',
            r'name[:\s]+([a-zA-Z\s,]+)',
            r'last\s*name[:\s]+([a-zA-Z\s]+)',
            r'first\s*name[:\s]+([a-zA-Z\s]+)',
        ]
        
        self.dob_patterns = [
            r'date\s*of\s*birth[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'dob[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'birth\s*date[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'born[:\s]+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
        ]
        
        self.mrn_patterns = [
            r'mrn[:\s]+(\w+)',
            r'medical\s*record\s*number[:\s]+(\w+)',
            r'patient\s*id[:\s]+(\w+)',
            r'id\s*number[:\s]+(\w+)',
        ]
        
        self.diagnosis_patterns = [
            r'diagnosis[:\s]+([^\.]+)',
            r'primary\s*diagnosis[:\s]+([^\.]+)',
            r'condition[:\s]+([^\.]+)',
        ]
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract patient information.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict[str, Any]: Processing results with patient data
        """
        start_time = datetime.now()
        
        try:
            # Check file format
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'extracted_text': '',
                    'patient_data': {},
                    'confidence_scores': {},
                    'processing_time': None,
                    'error_message': f'Unsupported file format: {file_extension}. Currently supported: {", ".join(self.supported_formats)}'
                }
            
            # Extract text based on file type
            if file_extension == '.docx':
                extracted_text = self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                extracted_text = self._extract_txt_text(file_path)
            else:
                extracted_text = "OCR processing not available in current deployment"
            
            # Clean the extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            # Extract patient information
            patient_data = self._extract_patient_info(cleaned_text)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                'success': True,
                'extracted_text': cleaned_text,
                'patient_data': patient_data,
                'confidence_scores': self._calculate_confidence_scores(patient_data, cleaned_text),
                'processing_time': processing_time,
                'error_message': None
            }
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {str(e)}")
            return {
                'success': False,
                'extracted_text': '',
                'patient_data': {},
                'confidence_scores': {},
                'processing_time': None,
                'error_message': str(e)
            }
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Text file extraction failed: {str(e)}")
            raise e
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise e
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\:\;\-\/\(\)]', '', text)
        
        # Normalize line breaks
        text = text.replace('\n\n', '\n').strip()
        
        return text
    
    def _extract_patient_info(self, text: str) -> Dict[str, Any]:
        """Extract structured patient information from text."""
        patient_data = {}
        
        # Convert to lowercase for pattern matching
        text_lower = text.lower()
        
        # Extract patient name
        name = self._extract_pattern(text_lower, self.name_patterns)
        if name:
            patient_data['name'] = name.title()
        
        # Extract date of birth
        dob = self._extract_pattern(text_lower, self.dob_patterns)
        if dob:
            patient_data['date_of_birth'] = self._normalize_date(dob)
        
        # Extract MRN
        mrn = self._extract_pattern(text_lower, self.mrn_patterns)
        if mrn:
            patient_data['medical_record_number'] = mrn.upper()
        
        # Extract diagnosis
        diagnosis = self._extract_pattern(text_lower, self.diagnosis_patterns)
        if diagnosis:
            patient_data['diagnosis'] = diagnosis.title()
        
        return patient_data
    
    def _extract_pattern(self, text: str, patterns: list) -> Optional[str]:
        """Extract text using regex patterns."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def _normalize_date(self, date_str: str) -> str:
        """Normalize date format."""
        # Simple date normalization
        date_str = date_str.replace('-', '/')
        return date_str
    
    def _calculate_confidence_scores(self, patient_data: Dict, text: str) -> Dict[str, float]:
        """Calculate confidence scores for extracted data."""
        scores = {}
        text_length = len(text)
        
        for field, value in patient_data.items():
            if value and text_length > 0:
                # Simple confidence based on text length and field presence
                scores[field] = min(0.85, 0.5 + (len(str(value)) / max(text_length, 100)))
            else:
                scores[field] = 0.0
        
        return scores